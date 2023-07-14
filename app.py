# from flask import Flask, request
# import pandas as pd
# import os
# from docx import Document

# app = Flask(__name__)

# @app.route('/envioExcel', methods=['POST'])
# def process_file():
#     file = request.files['file']
#     filename = file.filename

#     # Verifica a extensão do arquivo
#     if filename.endswith('.csv'):
#         df = pd.read_csv(file)
#     elif filename.endswith('.xlsx'):
#         df = pd.read_excel(file)
#     else:
#         return 'Formato de arquivo inválido. Por favor, envie um arquivo .csv ou .xlsx.'

#     template_dir = 'tC'
#     output_dir = 'cP'
#     doc_filename = 'dC.docx'

#     # Verifica se os diretórios existem
#     if not os.path.exists(template_dir):
#         return 'Diretório templateCertificado não encontrado.'

#     if not os.path.exists(output_dir):
#         os.makedirs(output_dir)

#     # Abre o arquivo de template
#     if not os.path.join(doc_filename):
#         return 'doc nao encontrado'
    
#     template_path = os.path.join(doc_filename)
#         # doc = Document(template_path)

#     doc = Document(template_path)
#         # return 'doc nao encontrado'

#     # Loop através das linhas do DataFrame
#     for index, row in df.iterrows():
#         nome = str(row['NOME'])
#         cpf = str(row['CPF'])

#         # Substitui os campos no documento
#         for p in doc.paragraphs:
#             if '{NOME}' in p.text:
#                 p.text = p.text.replace('{NOME}', nome)
#             if '{CPF}' in p.text:
#                 p.text = p.text.replace('{CPF}', cpf)

#         # Salva o documento atualizado em formato PDF
#         output_filename = f'certificado_{index+1}.pdf'
#         output_path = os.path.join(output_filename)
#         doc.save(output_path)

#     return 'Certificados gerados com sucesso!'

# if __name__ == '__main__':
#      app.run()









from flask import Flask, request, jsonify
import os
import pandas as pd
import docx
from docx2pdf import convert

app = Flask(__name__)

@app.route("/processar-arquivo", methods=["POST"])
def processar_arquivo():
    # Verificar se o arquivo foi enviado
    if "file" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"})

    file = request.files["file"]

    # Verificar a extensão do arquivo
    filename = file.filename
    file_ext = os.path.splitext(filename)[1]
    if file_ext not in [".csv", ".xlsx"]:
        return jsonify({"error": "Formato de arquivo inválido. Apenas arquivos CSV e XLSX são suportados"})

    # Salvar o arquivo em disco
    file.save(filename)

    # Processar o arquivo e extrair os nomes
    df = pd.read_csv(filename) if file_ext == ".csv" else pd.read_excel(filename)
    names = df["NOME"].tolist()
    cpfs = df["CPF"].astype(str).tolist()

    # Abrir o arquivo dC.docx e substituir os campos
    doc = docx.Document("tC/dC.docx")
    for i in range(len(names)):
        doc_new = docx.Document("tC/dC.docx")
        # j = 1
        for paragraph in doc_new.paragraphs:
            if "{NOME}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{NOME}", names[i])
            if "{CPF}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{CPF}", cpfs[i])
            # if "{NOME1}" in paragraph.text:
            #     paragraph.text = paragraph.text.replace("{NOME1}", names[j])
            # if "{CPF1}" in paragraph.text:
            #     paragraph.text = paragraph.text.replace("{CPF1}", cpfs[j])
        doc_new.save(f"CertificadoPronto/Certificado_{names[i]+'-'+cpfs[i]}.docx")

    # Converter os arquivos para PDF
    for i in range(len(names)):
        convert(f"CertificadoPronto/Certificado_{names[i]+'-'+names[i+1]}.docx")

    # Excluir o arquivo original
    # for i in range(len(names)):
    #     os.remove(f"CertificadoPronto/Certificado_{names[i]+'-'+cpfs[i]}.docx")

    return jsonify({"names": names})

if __name__ == "__main__":
    app.run()












